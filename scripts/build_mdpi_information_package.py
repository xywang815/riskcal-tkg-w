#!/usr/bin/env python3
"""Build the MDPI Information upload bundle and author-reference package."""

from __future__ import annotations

import argparse
import hashlib
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ACCENT = "245B78"
INK = "24313A"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def set_font(run, name="Calibri", east_asia="PingFang SC") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def add_bottom_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), ACCENT)
    borders.append(bottom)
    p_pr.append(borders)


def build_cover_letter(source: Path, output: Path) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(7)

    heading = document.add_paragraph()
    run = heading.add_run("COVER LETTER")
    set_font(run)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string(ACCENT)
    add_bottom_border(heading)

    for line in lines:
        stripped = line.strip()
        if not stripped or stripped == "# Cover Letter":
            continue
        paragraph = document.add_paragraph()
        if stripped.startswith("**") and stripped.endswith("**"):
            stripped = stripped[2:-2]
            bold = True
        else:
            bold = False
        stripped = stripped.replace("**", "").replace("*", "").replace("`", "")
        run = paragraph.add_run(stripped)
        set_font(run)
        run.bold = bold
        if stripped == "Sincerely,":
            paragraph.paragraph_format.space_before = Pt(8)
        if stripped.startswith("2 September"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    document.core_properties.title = "Cover Letter - Prequential Answer-Set Calibration"
    document.core_properties.author = "Xinyu Wang"
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def copy_file(source: Path, destination: Path) -> None:
    destination.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, destination)


def zip_tree(source_dir: Path, output_zip: Path) -> None:
    output_zip.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(output_zip, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as archive:
        for path in sorted(source_dir.rglob("*")):
            if path.is_file():
                archive.write(path, path.relative_to(source_dir))


def build(output_dir: Path, chinese_docx: Path) -> Path:
    if output_dir.exists():
        shutil.rmtree(output_dir)
    upload = output_dir / "upload_ready"
    source = output_dir / "manuscript_source"
    author = output_dir / "author_reference"
    for directory in (upload, source, author):
        directory.mkdir(parents=True, exist_ok=True)

    copy_file(ROOT / "paper/manuscript_mdpi_information.pdf", upload / "manuscript.pdf")
    copy_file(ROOT / "paper/manuscript_mdpi_information.tex", source / "manuscript.tex")
    copy_file(ROOT / "paper/references.bib", source / "references.bib")
    copy_file(ROOT / "LICENSE", source / "LICENSE")
    shutil.copytree(ROOT / "paper/Definitions", source / "Definitions")
    for figure in sorted((ROOT / "paper/figures/expansion").glob("expansion_*")):
        if figure.suffix.lower() in {".pdf", ".png", ".json"}:
            copy_file(figure, source / "figures/expansion" / figure.name)
    copy_file(
        ROOT / "paper/figures/expansion/SHA256SUMS.txt",
        source / "figures/expansion/SHA256SUMS.txt",
    )

    build_cover_letter(
        ROOT / "submission/information/COVER_LETTER_INFORMATION_20260902.md",
        upload / "cover_letter.docx",
    )
    zip_tree(source, upload / "manuscript_source.zip")
    shutil.rmtree(source)

    copy_file(chinese_docx, author / chinese_docx.name)
    for name in (
        "EXPANSION_RESULT_SUMMARY_20260902.md",
        "EXPANSION_CLAIM_LEDGER_20260901.md",
    ):
        copy_file(ROOT / "docs" / name, author / name)
    audit = ROOT / "paper/data/expansion/provenance/RELEASE_AUDIT_POST_MANUSCRIPT.json"
    if audit.exists():
        copy_file(audit, author / audit.name)

    copy_file(
        ROOT / "submission/information/PACKAGE_README_20260902.md",
        output_dir / "README.md",
    )
    copy_file(
        ROOT / "submission/information/SUBMISSION_CHECKLIST_20260902.md",
        output_dir / "SUBMISSION_CHECKLIST.md",
    )
    copy_file(
        ROOT / "submission/information/RELEASE_NOTES_v1.0.0.md",
        output_dir / "RELEASE_NOTES.md",
    )

    checksum_path = output_dir / "PACKAGE_SHA256SUMS.txt"
    records = []
    for path in sorted(output_dir.rglob("*")):
        if path.is_file() and path != checksum_path:
            records.append(f"{sha256(path)}  {path.relative_to(output_dir)}")
    checksum_path.write_text("\n".join(records) + "\n", encoding="utf-8")

    final_zip = output_dir.parent / f"{output_dir.name}.zip"
    if final_zip.exists():
        final_zip.unlink()
    zip_tree(output_dir, final_zip)
    return final_zip


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=ROOT / "submission/information/MDPI_Information_Submission_20260902",
    )
    parser.add_argument(
        "--chinese-docx",
        type=Path,
        default=ROOT / "submission/information/final_20260902/RiskCal-TKG_中文逐章翻译与解读.docx",
    )
    args = parser.parse_args()
    archive = build(args.output_dir.resolve(), args.chinese_docx.resolve())
    print(args.output_dir.resolve())
    print(archive)


if __name__ == "__main__":
    main()
