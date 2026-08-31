from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_TRANSLATION = Path(
    "/Users/wangxinyu/Desktop/RiskCal-TKG_MDPI版论文完整逐字翻译.docx"
)
OUTPUT = Path(
    "/Users/wangxinyu/Desktop/"
    "RiskCal-TKG_MDPI_Information_最新版中文逐字翻译稿_20260831.docx"
)
FIGURES = ROOT / "paper" / "figures" / "iclr_revision"

INK = RGBColor(11, 37, 69)
HEADING = RGBColor(46, 116, 181)
HEADING_DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(85, 85, 85)
LIGHT_FILL = "E8EEF5"
BODY_FONT = "Arial Unicode MS"
SANS_FONT = "Arial Unicode MS"


def set_east_asia_font(run, name: str) -> None:
    run.font.name = name
    props = run._element.get_or_add_rPr()
    fonts = props.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        props.insert(0, fonts)
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), name)


def style_run(run, size: float = 11, bold: bool = False, color=None, font=BODY_FONT) -> None:
    set_east_asia_font(run, font)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for node in list(grid):
        grid.remove(node)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    style_run(run, size=9, color=MUTED, font=SANS_FONT)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.33
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    body = doc.styles["Body Text"]
    body.base_style = normal
    body.font.name = BODY_FONT
    body.font.size = Pt(11)
    body._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    body.paragraph_format.space_after = Pt(8)
    body.paragraph_format.line_spacing = 1.33
    body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size, color, before, after in (
        ("Heading 1", 16, HEADING, 18, 10),
        ("Heading 2", 13, HEADING, 12, 6),
        ("Heading 3", 12, HEADING_DARK, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = SANS_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:eastAsia"), SANS_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = BODY_FONT
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    styles = doc.styles
    if "Equation CN" not in styles:
        eq = styles.add_style("Equation CN", 1)
    else:
        eq = styles["Equation CN"]
    eq.font.name = "Cambria Math"
    eq.font.size = Pt(10.5)
    eq._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    eq.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq.paragraph_format.space_before = Pt(4)
    eq.paragraph_format.space_after = Pt(6)
    eq.paragraph_format.keep_together = True

    if "Caption CN" not in styles:
        cap = styles.add_style("Caption CN", 1)
    else:
        cap = styles["Caption CN"]
    cap.font.name = BODY_FONT
    cap.font.size = Pt(9.5)
    cap.font.italic = False
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(10)
    cap.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("RiskCal-TKG | MDPI Information 中文逐字翻译稿")
    style_run(run, size=9, color=MUTED, font=SANS_FONT)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(footer)


def add_body(doc: Document, text: str, style: str = "Body Text"):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    style_run(run)
    return paragraph


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    style_run(run, size={1: 16, 2: 13, 3: 12}[level], bold=True,
              color=HEADING if level < 3 else HEADING_DARK, font=SANS_FONT)
    return paragraph


def add_equation(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="Equation CN")
    run = paragraph.add_run(text)
    style_run(run, size=10.5, font="Cambria Math")
    return paragraph


def add_caption(doc: Document, label: str, text: str):
    paragraph = doc.add_paragraph(style="Caption CN")
    first = paragraph.add_run(f"{label} ")
    style_run(first, size=9.5, bold=True)
    rest = paragraph.add_run(text)
    style_run(rest, size=9.5)
    return paragraph


def looks_like_equation(text: str) -> bool:
    markers = (" = ", " <= ", " >= ", "<=>", "subseteq", "sum_", "Pr{", "inf {", "O(B")
    return any(marker in text for marker in markers) and len(text) < 260


def normalize_reused_translation(text: str) -> str:
    replacements = {
        "预序列": "预序",
        "保形预测": "共形预测",
        "保形化": "共形化",
        "不符合度": "非一致性",
        "MVP weighted rule": "所实现的未修正加权规则",
        "主张边界": "结论边界",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def copy_reused_paragraph(doc: Document, source_paragraph) -> None:
    text = normalize_reused_translation(source_paragraph.text.strip())
    if not text:
        return
    style_name = source_paragraph.style.name
    if style_name == "Heading 1":
        add_heading(doc, text, 1)
        return
    if style_name == "Heading 2":
        add_heading(doc, text, 2)
        return
    if style_name == "List Bullet":
        add_body(doc, text, "List Bullet")
        return
    if looks_like_equation(text):
        add_equation(doc, text)
        return
    if text.startswith("命题 1") or text.startswith("命题 2") or text.startswith("推论"):
        paragraph = doc.add_paragraph(style="Body Text")
        label, remainder = text.split("。", 1) if "。" in text else (text, "")
        first = paragraph.add_run(label + ("。" if remainder else ""))
        style_run(first, bold=True)
        if remainder:
            rest = paragraph.add_run(remainder)
            style_run(rest)
        return
    add_body(doc, text)


def add_figure(doc: Document, image_name: str, label: str, caption: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(FIGURES / image_name), width=Inches(6.45))
    add_caption(doc, label, caption)


def add_appendix_tables(doc: Document) -> None:
    add_caption(
        doc,
        "表 A1.",
        "30% 删除率下基于历史的排序基线。仅训练基线只使用保留下来的训练事实；预序基线还使用首个测试时间戳之前的校准事实，并在每个完整测试时间戳批次之后更新历史。数值为五个随机种子的均值。",
    )
    rows = [
        ["方法", "历史", "MRR", "H@1", "H@10"],
        ["DistMult", "冻结评分器", "0.3105", "0.2310", "0.4661"],
        ["Frequency", "仅训练", "0.0925", "0.0437", "0.1845"],
        ["Relation-frequency", "仅训练", "0.1403", "0.0722", "0.2768"],
        ["Repeat", "仅训练", "0.2774", "0.2101", "0.4031"],
        ["Frequency", "预序", "0.0939", "0.0437", "0.1936"],
        ["Relation-frequency", "预序", "0.1532", "0.0812", "0.2998"],
        ["Repeat", "预序", "0.3631", "0.2827", "0.5114"],
    ]
    table = doc.add_table(rows=len(rows), cols=5)
    table.style = "Table Grid"
    set_table_geometry(table, [2300, 1900, 1720, 1720, 1720])
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx < 2 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            style_run(run, size=9.3, bold=r_idx == 0, font=SANS_FONT if r_idx == 0 else BODY_FONT)
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_FILL)

    add_caption(
        doc,
        "表 A2.",
        "30% 删除率下配对共享块 bootstrap 的块长敏感性。数值为 95% 置信区间；所有设置均使用 20,000 次重复和相同的五个随机种子。",
    )
    rows2 = [
        ["块长", "欠覆盖减少", "Micro 覆盖增益", "MRR 增益"],
        ["3", "[0.0580, 0.0877]", "[0.0631, 0.0913]", "[0.2082, 0.2285]"],
        ["7", "[0.0535, 0.0938]", "[0.0595, 0.0962]", "[0.2070, 0.2297]"],
        ["14", "[0.0492, 0.0987]", "[0.0557, 0.1012]", "[0.2083, 0.2281]"],
        ["21", "[0.0474, 0.0991]", "[0.0545, 0.1019]", "[0.2080, 0.2278]"],
    ]
    table2 = doc.add_table(rows=len(rows2), cols=4)
    table2.style = "Table Grid"
    set_table_geometry(table2, [1500, 2620, 2620, 2620])
    for r_idx, row in enumerate(rows2):
        for c_idx, value in enumerate(row):
            cell = table2.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            style_run(run, size=9.1, bold=r_idx == 0, font=SANS_FONT if r_idx == 0 else BODY_FONT)
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_FILL)


def build() -> None:
    old = Document(SOURCE_TRANSLATION)
    doc = Document()
    configure_document(doc)

    # Editorial-cover opening block.
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(72)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("MDPI INFORMATION 投稿稿 · 中文逐字翻译")
    style_run(run, size=11, bold=True, color=HEADING, font=SANS_FONT)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("时间知识图谱预测的预序答案集校准")
    style_run(run, size=26, bold=True, color=INK, font=SANS_FONT)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Prequential Answer-Set Calibration for Temporal Knowledge Graph Forecasting"
    )
    style_run(run, size=12, color=MUTED, font="Arial")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(24)
    run = meta.add_run(
        "作者：Xinyu Wang（王新宇）\n"
        "单位：安徽信息工程学院，中国芜湖 241000\n"
        "邮箱：xywang68@iflytek.com\n"
        "翻译基准：2026-08-31 官方 MDPI Information 模板英文稿"
    )
    style_run(run, size=11, color=MUTED)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(40)
    run = note.add_run(
        "翻译说明：正文按最新版英文稿逐段对应；公式符号、统计量和参考文献编号保持不变。"
    )
    style_run(run, size=9.5, color=MUTED)
    note.add_run().add_break(WD_BREAK.PAGE)

    add_heading(doc, "摘要", 1)
    add_body(
        doc,
        "时间知识图谱利用不断演化的关系证据支持事件预测，但标准排序指标无法说明模型能否返回既紧凑又具有可靠覆盖率的答案集合。现有校准研究大多假设数据静态或可交换。只使用过去时间戳批次更新的阈值，能否在时间漂移下修复可靠性，同时避免答案集合变得大到难以使用？我们在 ICEWS14 上采用严格防泄漏的预序协议研究这一问题。连续时间 DistMult 评分器在训练后被冻结，每个测试时间戳批次都必须先完成预测，随后才揭示其标签。在五个随机种子和最高 30% 的训练事实删除下，静态校准在目标覆盖率 0.90 时仅覆盖 0.8235 的观测标签。滚动校准达到 0.8998；其配对时间戳块欠覆盖减少量为 0.0723（95% 置信区间 [0.0535, 0.0938]）。这种修复代价较高：平均集合大小为 3,802.5，而实体总数为 7,128；36.5% 的唯一查询集合包含整个实体词表。仅在校准数据上从固定网格中选择的 RAPS 达到 0.8987 的覆盖率，平均集合大小为 3,319.0，但完整答案集合查询覆盖率降至 0.8874。因此，近期历史校准能够修复观测标签层面的可靠性，同时也暴露出显著的效用限制和子组覆盖限制。",
    )
    p = doc.add_paragraph(style="Body Text")
    lead = p.add_run("关键词：")
    style_run(lead, bold=True)
    rest = p.add_run(
        "时间知识图谱；链接预测；不确定性量化；预序评估；校准；共形预测；分布漂移；选择性预测"
    )
    style_run(rest)

    # Sections 1-4 from the earlier literal translation, limited to text that
    # is unchanged in the final English source.
    for index in range(10, 116):
        copy_reused_paragraph(doc, old.paragraphs[index])

    add_caption(doc, "算法 1.", "按时间戳成批的预序校准")
    steps = [
        "在训练时间戳上拟合 TKG 评分器，并冻结选定的检查点。",
        "在彼此不相交的测试前区间上，拟合可选漂移特征和选择器组件。",
        "使用最终初始校准区间初始化校准池 P。",
        "对于每个测试时间戳 t，在不使用当前标签的情况下对所有主体和客体查询打分；从严格过去的分数计算 q_t；然后输出全部预测集合。",
        "只有在第 4 步完成之后才揭示整个时间戳批次，将其非一致性分数加入 P，并继续处理下一个时间戳。",
    ]
    for step in steps:
        add_body(doc, step, "List Number")
    add_caption(
        doc,
        "图 1.",
        "受泄漏控制的评估协议。可选变体只改变如何从 P 估计阈值；任何方法都不能使用当前正在预测的时间戳标签。",
    )
    for index in range(128, 134):
        copy_reused_paragraph(doc, old.paragraphs[index])

    # Final Section 5.
    add_heading(doc, "5. 实验设计", 1)
    add_heading(doc, "5.1 数据与时间角色", 2)
    add_body(
        doc,
        "我们把 ICEWS14 作为受控事件预测案例。经过确定性标准化后，该数据集包含 7,128 个实体、230 种关系、90,730 条事实和 365 个按天计的时间戳。前 219 个时间戳用于训练评分器；时间戳 220–292 承担彼此分离的验证与校准角色；时间戳 293–365 构成预序测试数据流。三个部分相应包含 52,993、18,439 和 19,298 条事实。实体词表在测试前封闭并固定。按照第 4 节定义的实体覆盖约束，以 ρ∈{0,0.1,0.2,0.3} 的比例删除训练事实；开发、校准和测试事实保持不变。",
    )
    add_heading(doc, "5.2 评分器、比较方法与指标", 2)
    add_body(
        doc,
        "冻结评分器为连续时间 DistMult。Frequency 和基于历史的 Repeat 用作排序参照。Static、Rolling-1000、Fixed Weighted-1000 和探索性半衰期选择器共用相同的全实体分数张量。Rank、APS 和仅用验证数据选择的 RAPS 滚动规则作为不同的候选短名单操作点单独评估。",
    )
    add_body(
        doc,
        "过滤 MRR 与 Hits@1/3/10 衡量排序性能。可靠性通过观测标签覆盖率、正欠覆盖 (0.90−Cov_hat)_+、完整答案集合查询覆盖率和部分答案召回率来评价。效用通过集合大小的均值、中位数和 P90，以及预测集合等于全部 7,128 个实体的比例来衡量。关系方向审计只保留在所有随机种子中合计至少有 250 个观测标签、且每个随机种子均有表示的组。该支持度筛选仅定义一个诊断子集，不构成条件覆盖保证。",
    )
    add_heading(doc, "5.3 确认性协议与统计推断", 2)
    add_body(
        doc,
        "主要设置为删除 30% 的训练历史，目标覆盖率为 0.90。主要估计量为：",
    )
    add_equation(
        doc,
        "Δ_roll = (1 / (|S||T|)) Σ_{s∈S} Σ_{t∈T} d_{s,t},",
    )
    add_equation(
        doc,
        "d_{s,t} = (0.90 − Cov_hat_Static^{s,t})_+ − (0.90 − Cov_hat_Rolling^{s,t})_+ .",
    )
    add_body(
        doc,
        "它衡量从 Static 到 Rolling-1000 的时间戳宏平均正欠覆盖减少量。评分器使用五个随机种子 {17,29,43,59,71}。我们报告配对的“随机种子重采样 + 圆形移动时间块 bootstrap”：重复 20,000 次，块长为 7 个时间戳，并在每次重复中为两种方法使用完全相同的时间戳块。块长 3、14 和 21 用于敏感性检查。代码、最终配置、预处理步骤和面向论文的派生表格见 https://github.com/xywang815/riskcal-tkg-w。",
    )

    # Final Section 6.
    add_heading(doc, "6. 结果", 1)
    add_body(
        doc,
        "全部 20 个“随机种子 × 删除率”条件均完成。图 2 汇总排序、校准、集合大小和关系方向结果；(b) 与 (c) 中的误差棒为五个随机种子的标准差。这里研究的是固定评分器的校准，而不是提出新的排序模型。",
    )
    add_figure(
        doc,
        "empirical_overview.png",
        "图 2.",
        "ICEWS14 主要证据。(a) 受控删除训练事实时的过滤 MRR；(b) Static 与 Rolling-1000 的观测标签覆盖率，虚线为 0.90 目标；(c) 按标签加权的平均预测集合大小；(d) 76 个满足支持条件的关系方向组中的最低覆盖率。",
    )
    add_heading(doc, "6.1 排序与预序可靠性", 2)
    add_body(
        doc,
        "在 30% 删除率下，连续时间 DistMult 的 MRR 为 0.3105，而 Frequency 为 0.0925。二者的配对共享块差异为 0.2180（95% 置信区间 [0.2070, 0.2297]，p<10^−4）。只使用训练历史的 Repeat 基线达到 MRR 0.2774；当校准历史和已经完成的测试批次可用后，预序 Repeat 达到 0.3631（见附录表 A1）。这说明 ICEWS14 具有很强的重复性，因此本文不把该评分器描述为最先进预测器。",
    )
    add_body(
        doc,
        "在 30% 删除率下，Static 只覆盖 0.8235 的观测标签。Rolling-1000 达到 0.8998，并把时间戳宏平均正欠覆盖减少 0.0723（95% 置信区间 [0.0535, 0.0938]，p<10^−4）。当块长取 3、7、14 和 21 时，置信区间下界始终为正（附录表 A2）。Fixed Weighted-1000 和选定的半衰期规则分别达到 0.8988 与 0.8989；这些几乎相同的数值不支持在短 Rolling-1000 池内近期加权具有优势。",
    )
    add_heading(doc, "6.2 查询语义与子组表现", 2)
    add_body(
        doc,
        "每个随机种子包含 33,964 个唯一主体/客体查询和 38,596 个观测标签。多答案查询占唯一查询的 10.1%，最多包含 17 个已记录答案。在 30% 删除率下，Rolling-1000 把完整答案集合查询覆盖率从 0.8267 提高到 0.9020，把部分答案召回率从 0.8463 提高到 0.9174。这种可靠性提升代价很高：按标签加权的平均集合大小从 2,748.6 增加到 3,802.5，P90 达到全部 7,128 个实体，并且 36.5% 的唯一查询预测集合包含完整实体词表。",
    )
    add_body(
        doc,
        "分组审计说明平均覆盖率并不足够。在 30% 删除率下，满足支持条件的关系方向组最低覆盖率从 Static 的 0.6274 提高到 Rolling-1000 的 0.7625；低于 0.90 的组比例从 63.2% 降到 22.4%。然而，最弱组仍远低于目标，因此该结果不是关系条件覆盖保证。",
    )
    add_heading(doc, "6.3 候选短名单与稳健性诊断", 2)
    add_body(
        doc,
        "图 3 比较候选短名单操作点、窗口选择和反馈延迟。Margin rolling 的观测标签覆盖率为 0.8998，完整答案集合查询覆盖率为 0.9020，唯一查询平均集合大小为 4,069.4。Rank rolling 把均值降至 3,668.2，并消除完整词表集合，但完整答案集合查询覆盖率降至 0.8881。APS rolling 的平均集合更大，为 4,911.8 个实体。",
    )
    add_body(
        doc,
        "仅用验证数据选择的 RAPS rolling 使用温度 T=1、选择容差 0.02，以及第 4 节规定的固定 21 候选网格。全部 20 个条件都选择 k=50、λ=10^−4。在 30% 删除率下，其观测标签覆盖率为 0.8987，集合大小的均值、中位数和 P90 分别为 3,319.0、2,924.6 和 4,646.5；完整答案集合查询覆盖率为 0.8874。因此，该规则改善了平均集合大小，但其尾部大小并不优于 rank rolling，也没有保持 margin rolling 的查询层覆盖率。",
    )
    add_figure(
        doc,
        "utility_robustness.png",
        "图 3.",
        "30% 删除率下的可靠性—效用与稳健性诊断。(a) 四种候选短名单评分规则；(b) Static、expanding、按分数计数窗口和按时间窗口校准；(c) 已完成时间戳批次在进入校准池前被暂扣时的 Rolling-1000 覆盖率。",
    )
    add_body(
        doc,
        "窗口审计把在线更新与近期性分离。Expanding 校准达到 0.8590 的覆盖率；Rolling-250、Rolling-500 和 Rolling-2000 分别达到 0.9129、0.9001 和 0.8964。默认 1,000 分数池在中位数上只跨越两个时间戳块，而 7 天半衰期的有效样本量为 997.6。因此，选择器在全部 20 个默认窗口条件中都选择等权。额外反馈延迟会单调降低 Rolling-1000：覆盖率从无延迟时的 0.8998 降到暂扣七个批次后的 0.8884。由此可见，可靠性修复依赖及时标签与校准窗口；这些因素没有被总平均结果掩盖。",
    )
    add_body(
        doc,
        "完整的五随机种子、四删除率实验在一张 RTX 4090 上耗时 88.4 分钟。每个条件的平均训练、校准和推理时间分别为 203.6、13.0 和 45.4 秒，峰值 CUDA 显存约为 1.05 GiB。这些数值只描述当前的 ICEWS14 全实体评估，不代表更大实体词表的开销。",
    )

    # Final merged Section 7.
    add_heading(doc, "7. 讨论与结论", 1)
    add_heading(doc, "7.1 主要发现与解释", 2)
    add_body(
        doc,
        "本研究把排序、可靠性和效用分开考察。尽管冻结评分器仍具有有用的排序能力，早期窗口 Static 校准在后续 ICEWS14 数据流上仍出现欠覆盖。只从严格过去的时间戳批次更新阈值，可以恢复接近目标的平均观测标签覆盖率。但这种收益并非没有代价：margin 集合往往大到难以直接检查，满足支持条件的关系方向组仍可能严重欠覆盖。Rank 与仅用验证数据选择的 RAPS rolling 给出更实际的操作点，但较小集合会降低完整答案集合查询覆盖率。因此，证据支持把近期历史再校准视为诊断性修复，而不支持某一种短名单规则全面占优，也不表示排序性能得到提升。",
    )
    add_body(
        doc,
        "命题 2 中的确定性分解解释了这一现象，同时没有声称新的覆盖定理。近期池可能减少时间错配，但有限池估计误差和并列分数原子仍然存在。在本实验中，显式半衰期加权几乎没有额外收益，因为默认分数窗口只跨越很少的时间戳。这个负结果同样有信息价值：当候选历史实际上难以区分时，更复杂的选择器不会带来帮助。",
    )
    add_heading(doc, "7.2 局限性与实际意义", 2)
    add_body(
        doc,
        "证据仅来自 ICEWS14 这一具有较强重复性的政治事件基准，以及一个可以精确枚举全部实体的 DistMult 评分器。受控删除是一种明确规定的压力测试，而不是事件缺失的生成模型。五个随机种子限制了对训练随机性的推断；移动块 bootstrap 仍假设其块方案能够表示局部时间依赖。ICEWS 中记录的事件可能不完整或含噪，因此覆盖率针对的是已记录标签，而不是所有真实世界事件。关系方向分析经过支持度筛选，不构成条件覆盖。最后，在任意时间漂移下，本文的经验 rolling、weighted 和 selected 规则都没有有限样本、分布无关保证。",
    )
    add_body(
        doc,
        "在部署中，实际决策单位应是一个完整操作点，而不是单一覆盖率数字。系统应同时报告观测标签覆盖率、完整答案集合查询覆盖率、尾部集合大小、反馈延迟和子组表现。若为更大实体词表引入近似检索，则必须在校准前审计候选召回率，因为遗漏真实实体会改变被校准的目标事件本身。",
    )
    add_heading(doc, "7.3 未来工作与结论", 2)
    add_body(
        doc,
        "后续研究应覆盖更多时间图、循环或 Transformer 型 TKG backbone、重复性较低的领域以及其他历史缺失机制。方法层面的方向包括：具有明确非可交换保证的校准器、关系感知评价，以及能够控制候选损失的检索感知短名单。还需要更长的校准历史，才能公平判断由特征驱动的半衰期选择。",
    )
    add_body(
        doc,
        "在本文规定的 ICEWS14 协议内，核心结论更窄并且完全是经验性的。在 30% 删除率下，Rolling-1000 把观测标签覆盖率从 0.8235 提高到 0.8998，配对欠覆盖减少量为 0.0723（95% 置信区间 [0.0535, 0.0938]）。这种修复伴随 3,802.5 的平均集合大小和明显的子组失败。仅用验证数据选择的 RAPS rolling 在覆盖率 0.8987 时，把唯一查询平均集合大小降至 3,319.0，但完整答案集合查询覆盖率降至 0.8874。这些结果构成可复现的可靠性—效用诊断；它们不构成新的分布无关覆盖保证。",
    )

    add_heading(doc, "作者贡献", 1)
    add_body(doc, normalize_reused_translation(old.paragraphs[234].text))
    add_heading(doc, "基金资助", 1)
    add_body(doc, "本研究由安徽省高校自然科学研究项目资助，项目编号 2023AH052916。")
    add_heading(doc, "机构审查委员会声明", 1)
    add_body(doc, "不适用。")
    add_heading(doc, "知情同意声明", 1)
    add_body(doc, "不适用。")
    add_heading(doc, "数据可用性声明", 1)
    add_body(
        doc,
        "实验使用公开的 ICEWS14 基准。源代码、实验配置、固定随机种子、面向论文的派生结果表和图生成材料可在 https://github.com/xywang815/riskcal-tkg-w 获取。原始基准文件应按照其分发条款从原始公开来源获得。该仓库提供本研究使用的预处理与标准化脚本。",
    )
    add_heading(doc, "利益冲突", 1)
    add_body(
        doc,
        "作者声明不存在利益冲突。资助方未参与研究设计、数据收集、分析或解释、论文写作以及是否发表结果的决定。",
    )

    add_heading(doc, "附录 A. 额外诊断", 1)
    add_appendix_tables(doc)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_heading(doc, "参考文献", 1)
    add_body(doc, "说明：书目信息保留英文原文，以便检索、投稿和 DOI 核对。")
    for index in range(252, 272):
        paragraph = doc.add_paragraph(style="Body Text")
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(5)
        run = paragraph.add_run(old.paragraphs[index].text.strip())
        style_run(run, size=9.5, font="Arial")

    core = doc.core_properties
    core.title = "时间知识图谱预测的预序答案集校准：中文逐字翻译稿"
    core.author = "Xinyu Wang"
    core.subject = "MDPI Information 投稿稿中文逐字翻译"
    core.keywords = "时间知识图谱, 预序校准, ICEWS14, 共形预测"
    core.comments = "Translated from the finalized 2026-08-31 English submission source."

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
