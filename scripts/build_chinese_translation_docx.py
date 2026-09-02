#!/usr/bin/env python3
"""Build the reader-facing Chinese translation from its tracked Markdown source."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ACCENT = "245B78"
LIGHT = "EAF1F5"
INK = "24313A"
MUTED = "5B6770"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, name: str = "Calibri", east_asia: str = "PingFang SC") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    set_run_font(run)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.92)
    section.right_margin = Inches(0.92)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.28
    normal.paragraph_format.space_after = Pt(5)

    for style_name, size, before, after, color in (
        ("Title", 22, 0, 14, ACCENT),
        ("Heading 1", 16, 18, 9, ACCENT),
        ("Heading 2", 13, 13, 6, ACCENT),
        ("Heading 3", 11.5, 9, 4, "304B5A"),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "RiskCal-TKG  |  中文逐章翻译与理解附录"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run)
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string(MUTED)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_horizontal_rule(document: Document, color: str = ACCENT) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(8)
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def add_inline(paragraph, text: str) -> None:
    token_re = re.compile(r"(\*\*.+?\*\*|`.+?`|https?://\S+)")
    position = 0
    for match in token_re.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            set_run_font(run)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            set_run_font(run)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, "Menlo", "PingFang SC")
            run.font.size = Pt(9.3)
            run.font.color.rgb = RGBColor.from_string("7A3E16")
        else:
            run = paragraph.add_run(token.rstrip(".,;，。；"))
            set_run_font(run)
            run.font.color.rgb = RGBColor.from_string("175C8C")
            run.underline = True
            trailing = token[len(token.rstrip(".,;，。；")) :]
            if trailing:
                tail = paragraph.add_run(trailing)
                set_run_font(tail)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run)


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("RISKCAL-TKG")
    set_run_font(run)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(ACCENT)

    for _ in range(2):
        document.add_paragraph()

    title = document.add_paragraph()
    title.style = document.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("时序知识图谱预测的\n预序贯答案集校准")
    set_run_font(run)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("英文投稿稿件的中文逐章翻译与理解附录")
    set_run_font(run)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    add_horizontal_rule(document)

    meta = document.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.autofit = False
    meta.columns[0].width = Inches(1.25)
    meta.columns[1].width = Inches(5.45)
    values = (
        ("英文题目", "Prequential Answer-Set Calibration for Temporal Knowledge Graph Forecasting"),
        ("作者", "Xinyu Wang"),
        ("单位", "安徽信息工程学院，中国安徽芜湖 241000"),
        ("说明", "忠实翻译正文；末尾理解附录不属于英文投稿稿件"),
    )
    for row, (label, value) in zip(meta.rows, values):
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run)
                    run.font.size = Pt(10)
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_shading(row.cells[0], LIGHT)

    document.add_paragraph()
    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(26)
    run = note.add_run("编制日期：2026 年 9 月 2 日")
    set_run_font(run)
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    document.add_page_break()


def parse_table(lines: list[str], start: int, document: Document) -> int:
    table_lines = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        table_lines.append(lines[index].strip())
        index += 1
    rows = [[cell.strip() for cell in line.strip("|").split("|")] for line in table_lines]
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in rows[1]):
        rows.pop(1)
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, ACCENT)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run)
                    run.font.size = Pt(9)
                    if r_idx == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
    document.add_paragraph().paragraph_format.space_after = Pt(1)
    return index


def build(markdown_path: Path, output_path: Path) -> None:
    document = Document()
    configure_document(document)
    add_cover(document)

    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    index = 1  # Cover already contains the first Markdown H1.
    paragraph_buffer: list[str] = []

    def flush_buffer() -> None:
        if not paragraph_buffer:
            return
        text = " ".join(item.strip() for item in paragraph_buffer).strip()
        paragraph_buffer.clear()
        if text:
            paragraph = document.add_paragraph()
            add_inline(paragraph, text)

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            flush_buffer()
            index += 1
            continue
        if stripped.startswith("|"):
            flush_buffer()
            index = parse_table(lines, index, document)
            continue
        image_match = re.fullmatch(r"!\[(.+?)\]\((.+?)\)", stripped)
        if image_match:
            flush_buffer()
            image_path = (markdown_path.parent / image_match.group(2)).resolve()
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_with_next = True
            run = paragraph.add_run()
            run.add_picture(str(image_path), width=Inches(6.35))
            caption = document.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.paragraph_format.space_after = Pt(8)
            cap_run = caption.add_run(image_match.group(1))
            set_run_font(cap_run)
            cap_run.italic = True
            cap_run.font.size = Pt(9)
            cap_run.font.color.rgb = RGBColor.from_string(MUTED)
            index += 1
            continue
        heading_match = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading_match:
            flush_buffer()
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            if level == 1 and text.startswith("理解附录"):
                document.add_page_break()
            paragraph = document.add_paragraph(style=f"Heading {level}")
            add_inline(paragraph, text)
            index += 1
            continue
        if stripped.startswith("> "):
            flush_buffer()
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.right_indent = Inches(0.15)
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(8)
            add_inline(paragraph, stripped[2:])
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor.from_string(MUTED)
                run.italic = True
            index += 1
            continue
        list_match = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if list_match:
            flush_buffer()
            paragraph = document.add_paragraph(style="List Number")
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.first_line_indent = Inches(-0.12)
            add_inline(paragraph, list_match.group(2))
            index += 1
            continue
        if stripped.startswith("- "):
            flush_buffer()
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.left_indent = Inches(0.25)
            add_inline(paragraph, stripped[2:])
            index += 1
            continue
        paragraph_buffer.append(stripped)
        index += 1

    flush_buffer()

    document.core_properties.title = "时序知识图谱预测的预序贯答案集校准：中文逐章翻译与理解附录"
    document.core_properties.author = "Xinyu Wang"
    document.core_properties.subject = "MDPI Information manuscript Chinese translation"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, default=Path("paper/manuscript_mdpi_information_zh.md"))
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("submission/information/final_20260902/RiskCal-TKG_中文逐章翻译与解读.docx"),
    )
    args = parser.parse_args()
    build(args.source.resolve(), args.output.resolve())
    print(args.output.resolve())


if __name__ == "__main__":
    main()
